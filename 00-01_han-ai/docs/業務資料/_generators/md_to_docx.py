# -*- coding: utf-8 -*-
"""業務資料フォルダ内の .md を、Google ドキュメント変換用に .docx へ変換する。
軽量な自前パーサ（見出し/太字/箇条書き(入れ子)/番号/表/コードブロック/引用/罫線）。"""
import os, re, glob
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = r"C:\Users\ooto\work\ClaudeCode\kurashi-no-dodai-log\00-01_han-ai\docs\業務資料"
JP = "Yu Gothic"
MONO = "Consolas"
INK = RGBColor(0x1A, 0x22, 0x33)
SUB = RGBColor(0x5B, 0x64, 0x72)
ACC = RGBColor(0x2F, 0x6F, 0x5E)

INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")


def set_run_font(run, name=JP, size=None, bold=False, color=None, italic=False):
    run.font.name = name
    run.font.bold = bold
    run.font.italic = italic
    if size:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    rPr = run._r.get_or_add_rPr()
    for tag in ("w:rFonts",):
        el = rPr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            rPr.append(el)
        for a in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            el.set(qn(a), name)


def add_inline(par, text, base_size=None, base_color=INK):
    """**bold** と `code` だけ解釈し、他は素のテキスト。"""
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = par.add_run(part[2:-2])
            set_run_font(r, JP, base_size, bold=True, color=base_color)
        elif part.startswith("`") and part.endswith("`"):
            r = par.add_run(part[1:-1])
            set_run_font(r, MONO, base_size, color=base_color)
        else:
            r = par.add_run(part)
            set_run_font(r, JP, base_size, color=base_color)


def shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hex_fill)
    tcPr.append(sh)


def convert(md_path):
    with open(md_path, encoding="utf-8") as f:
        lines = f.read().splitlines()
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = JP
    st.font.size = Pt(10.5)
    st.font.color.rgb = INK
    st.element.rPr.rFonts.set(qn("w:eastAsia"), JP)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]

        # code block
        if line.strip().startswith("```"):
            i += 1
            buf = []
            while i < n and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)
            pPr = p._p.get_or_add_pPr()
            sh = OxmlElement("w:shd")
            sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), "F2F7F5")
            pPr.append(sh)
            for k, cl in enumerate(buf):
                r = p.add_run(cl)
                set_run_font(r, MONO, 9, color=INK)
                if k != len(buf) - 1:
                    r.add_break()
            continue

        # table block
        if line.lstrip().startswith("|") and i + 1 < n and re.match(r"^\s*\|[\s:|-]+\|\s*$", lines[i + 1]):
            rows = []
            while i < n and lines[i].lstrip().startswith("|"):
                rows.append(lines[i])
                i += 1
            cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
            cells = [cells[0]] + cells[2:]  # drop separator
            ncol = max(len(r) for r in cells)
            tbl = doc.add_table(rows=len(cells), cols=ncol)
            tbl.style = "Table Grid"
            tbl.autofit = True
            for ri, row in enumerate(cells):
                for ci in range(ncol):
                    cell = tbl.cell(ri, ci)
                    cell.text = ""
                    par = cell.paragraphs[0]
                    val = row[ci] if ci < len(row) else ""
                    add_inline(par, val, base_size=9.5)
                    if ri == 0:
                        shade(cell, "E4EFEA")
                        for rr in par.runs:
                            rr.font.bold = True
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            continue

        # headings
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            lvl = len(m.group(1))
            h = doc.add_heading(level=min(lvl, 4))
            h.text = ""
            add_inline(h, m.group(2),
                       base_size={1: 18, 2: 14, 3: 12, 4: 11}.get(lvl, 11),
                       base_color=ACC if lvl <= 2 else INK)
            for r in h.runs:
                r.font.bold = True
            i += 1
            continue

        # horizontal rule
        if re.match(r"^\s*---+\s*$", line):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            bd = OxmlElement("w:pBdr")
            bot = OxmlElement("w:bottom")
            bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6")
            bot.set(qn("w:space"), "1"); bot.set(qn("w:color"), "C9D2CE")
            bd.append(bot); pPr.append(bd)
            i += 1
            continue

        # blockquote
        if line.lstrip().startswith(">"):
            buf = []
            while i < n and lines[i].lstrip().startswith(">"):
                buf.append(re.sub(r"^\s*>\s?", "", lines[i]))
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            add_inline(p, " ".join(x for x in buf if x.strip()), base_size=10, base_color=SUB)
            continue

        # list item (bullet / numbered), with nesting by indent
        m = re.match(r"^(\s*)([-*]|\d+\.)\s+(.*)$", line)
        if m:
            indent = len(m.group(1))
            ordered = bool(re.match(r"\d+\.", m.group(2)))
            style = "List Number" if ordered else "List Bullet"
            lvl = min(indent // 2, 2)
            if lvl and not ordered:
                style = "List Bullet 2" if lvl == 1 else "List Bullet 3"
            try:
                p = doc.add_paragraph(style=style)
            except KeyError:
                p = doc.add_paragraph(style="List Bullet")
            # checkbox items
            txt = m.group(3)
            txt = re.sub(r"^\[[ xX]\]\s*", "☐ ", txt)
            add_inline(p, txt, base_size=10.5)
            i += 1
            continue

        # blank
        if not line.strip():
            i += 1
            continue

        # plain paragraph
        p = doc.add_paragraph()
        add_inline(p, line, base_size=10.5)
        i += 1

    out = os.path.splitext(md_path)[0] + ".docx"
    doc.save(out)
    return out


targets = []
for md in sorted(glob.glob(os.path.join(ROOT, "**", "*.md"), recursive=True)):
    base = os.path.basename(md)
    if base.startswith("_"):        # 04_DX の未整備プレースホルダは対象外
        continue
    targets.append(md)

for md in targets:
    out = convert(md)
    d = Document(out)
    print("OK  {}  (paras={}, tables={})".format(
        os.path.relpath(out, ROOT).replace(os.sep, "/"),
        len(d.paragraphs), len(d.tables)))
print("done:", len(targets), "docx")
